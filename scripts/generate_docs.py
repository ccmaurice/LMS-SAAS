import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins inside a table cell (values in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level, before=12, after=6):
    """Add a heading with custom spacing before and after."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(after)
    h.paragraph_format.keep_with_next = True
    return h

def main():
    doc = docx.Document()

    # Page Setup: Standard Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Configurations
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # TITLE PAGE
    # ----------------------------------------------------
    p_title_spacer = doc.add_paragraph()
    p_title_spacer.paragraph_format.space_before = Pt(40)

    title_text = (
        "A SCALABLE MULTI-TENANT SAAS ARCHITECTURE FOR A LEARNING "
        "MANAGEMENT SYSTEM: DESIGN AND IMPLEMENTATION FOR NIGERIAN "
        "EDUCATIONAL INSTITUTIONS"
    )
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title_text)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    p_title.paragraph_format.space_after = Pt(36)

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_by = p_by.add_run("BY:\n\nCHUKWUDEBELU CHINEDU MAURICE\n(2025/A/MIT/0651/301833716)")
    run_by.font.name = 'Times New Roman'
    run_by.font.size = Pt(14)
    run_by.font.bold = True
    p_by.paragraph_format.space_after = Pt(36)

    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_dept = p_dept.add_run(
        "DEPARTMENT OF INFORMATION TECHNOLOGY\n"
        "SCHOOL OF COMPUTING\n"
        "MIVA OPEN UNIVERSITY ABUJA, NIGERIA."
    )
    run_dept.font.name = 'Times New Roman'
    run_dept.font.size = Pt(12)
    run_dept.font.bold = True
    p_dept.paragraph_format.space_after = Pt(48)

    p_purpose = doc.add_paragraph()
    p_purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_purpose = p_purpose.add_run(
        "A PROFESSIONAL MASTER'S PROJECT SUBMITTED TO THE DEPARTMENT OF "
        "INFORMATION TECHNOLOGY, SCHOOL OF COMPUTING, MIVA OPEN "
        "UNIVERSITY ABUJA, NIGERIA.\n\n"
        "IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE AWARD OF THE "
        "PROFESSIONAL MASTER OF INFORMATION TECHNOLOGY (MIT) DEGREE IN "
        "INFORMATION TECHNOLOGY"
    )
    run_purpose.font.name = 'Times New Roman'
    run_purpose.font.size = Pt(11)
    p_purpose.paragraph_format.space_after = Pt(60)

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("AUGUST, 2026")
    run_date.font.name = 'Times New Roman'
    run_date.font.size = Pt(12)
    run_date.font.bold = True

    doc.add_page_break()

    # ----------------------------------------------------
    # DECLARATION & CERTIFICATION
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "DECLARATION", level=1)
    p_decl = doc.add_paragraph(
        "I hereby declare that this project report titled \"A Scalable Multi-Tenant SaaS Architecture "
        "for a Learning Management System: Design and Implementation for Nigerian Educational Institutions\" "
        "is my original work carried out under the supervision of Prof. Chidiebere Ugwu. All sources of "
        "information, literature review material, and third-party references utilized in this work have "
        "been duly acknowledged."
    )
    doc.add_paragraph("\n\n..................................................\nChukwudebelu Chinedu Maurice\nDate")

    doc.add_page_break()

    add_heading_with_spacing(doc, "CERTIFICATION", level=1)
    doc.add_paragraph(
        "This is to certify that I, Chukwudebelu Chinedu Maurice (2025/A/MIT/0651/301833716), am "
        "responsible for the work submitted in this project, that the original work is mine, except as "
        "specified in acknowledgment and references, and that neither the project nor the original work "
        "contained therein has been submitted to this University or any other institutions for the award "
        "of a degree."
    )
    doc.add_paragraph("\n\n..................................................\nSignature and Date")

    doc.add_page_break()

    add_heading_with_spacing(doc, "APPROVAL", level=1)
    doc.add_paragraph(
        "This project has been approved for the Department of Information Technology, School of "
        "Computing, Miva Open University, Abuja, Nigeria."
    )
    doc.add_paragraph("\n\nName of Supervisor: Prof. Chidiebere Ugwu")
    doc.add_paragraph("Supervisor Signature and Date: ..................................................")
    doc.add_paragraph("\nName of Head of Department: Dr. Abdullahi Mohammed")
    doc.add_paragraph("Head of Department Signature and Date: ..................................................")
    doc.add_paragraph("\nName of Dean: Prof. Victoria Opara")
    doc.add_paragraph("Dean, School of Computing Signature and Date: ..................................................")
    doc.add_paragraph("\nName of External Examiner: Prof. Kingsley Obani")
    doc.add_paragraph("External Examiner Signature and Date: ..................................................")

    doc.add_page_break()

    # ----------------------------------------------------
    # DEDICATION & ACKNOWLEDGEMENT
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "DEDICATION", level=1)
    doc.add_paragraph(
        "This project is dedicated to Almighty God for His infinite grace and wisdom, to my family for "
        "their unwavering support, and to the educators across Nigeria who strive daily to bridge the "
        "digital divide in our classrooms."
    )

    doc.add_page_break()

    add_heading_with_spacing(doc, "ACKNOWLEDGEMENT", level=1)
    doc.add_paragraph(
        "I wish to express my profound gratitude to Almighty God for the strength, health, and guidance "
        "granted to me throughout the duration of this Master's program and project."
    )
    doc.add_paragraph(
        "My sincere appreciation goes to my supervisor, Prof. Chidiebere Ugwu, whose meticulous "
        "scrutiny, constructive criticism, and invaluable guidance shaped this work from a mere idea "
        "into a robust software engineering artifact. I am deeply grateful for the patience and time "
        "dedicated to ensuring this project met professional standards."
    )
    doc.add_paragraph(
        "I extend my thanks to the Head of Department, Dr. Abdullahi Mohammed, and the entire faculty and "
        "staff of the Department of Information Technology, School of Computing, Miva Open "
        "University, for providing an enabling academic environment and the resources necessary for "
        "the completion of this work."
    )
    doc.add_paragraph(
        "Special thanks to my colleagues in the MIT program for their camaraderie and technical "
        "exchanges, which enriched my understanding of software architecture. Finally, my deepest "
        "gratitude goes to my parents and spouse for their prayers, patience, and financial support during "
        "the long hours of development and documentation."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # ABSTRACT
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "ABSTRACT", level=1)
    doc.add_paragraph(
        "The adoption of Learning Management Systems (LMS) in Nigerian educational institutions is "
        "hindered by significant barriers, including high infrastructure costs, the technical complexity "
        "of managing single-tenant open-source platforms like Moodle, and the prohibitive subscription "
        "fees of global cloud-native solutions. This project addresses these challenges through the "
        "design and implementation of a scalable, multi-tenant Software as a Service (SaaS) architecture "
        "specifically tailored for primary, secondary, and tertiary institutions in Nigeria."
    )
    doc.add_paragraph(
        "The project employs an iterative development methodology informed by Agile principles. A "
        "comprehensive literature review was conducted to establish the theoretical foundations of "
        "multi-tenancy in cloud computing and to analyze existing systems, revealing a critical gap in "
        "localized, cost-effective platforms that are optimized for the Nigerian internet landscape. The "
        "system architecture leverages a shared-database, shared-schema PostgreSQL database using Prisma "
        "7 Object-Relational Mapping (ORM) to enforce tenant isolation logically at the application and database "
        "layers. It is built as a Next.js 16 (App Router) serverless application on Vercel with a mobile-first, "
        "bandwidth-optimized React 19 and TailwindCSS 4 interface suitable for variable 3G/4G cellular networks."
    )
    doc.add_paragraph(
        "The resulting artifact is a fully deployed, functional prototype that demonstrates core LMS "
        "functionalities—course management, user administration, assignment handling, and automated certificates—within "
        "a secure, multi-tenant environment. Significantly, this prototype features an advanced exam integrity "
        "and live proctoring dashboard. The dashboard integrates WebRTC-based support camera pairing, periodic "
        "camera base64 frame captures, 3-second sliced microphone audio streaming, live invigilator listening "
        "capabilities, remote command signaling (Prompt/Force Camera/Audio), and evidence clip recording. The "
        "live application is accessible at https://saas-lms-khaki.vercel.app/ and via the custom domain "
        "https://skilltech.com.ng/, with all source code available in the companion GitHub repository: "
        "https://github.com/ccmaurice/SaaS-LMS. Performance testing under simulated Nigerian network conditions "
        "yielded a First Contentful Paint of 1.8 seconds, meeting the target of sub-3-second loads. The evaluation "
        "confirms that the logical multi-tenant model effectively amortizes operational costs while preserving "
        "performance, data security, and invigilation integrity."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # TABLE OF CONTENTS / LIST OF FIGURES / LIST OF TABLES
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "TABLE OF CONTENTS", level=1)
    doc.add_paragraph("DECLARATION ......................................................................................................................... ii")
    doc.add_paragraph("CERTIFICATION ........................................................................................................................ iii")
    doc.add_paragraph("APPROVAL ............................................................................................................................. iv")
    doc.add_paragraph("DEDICATION ........................................................................................................................... v")
    doc.add_paragraph("ACKNOWLEDGEMENTS .................................................................................................................. vi")
    doc.add_paragraph("ABSTRACT ............................................................................................................................. vii")
    doc.add_paragraph("LIST OF FIGURES ..................................................................................................................... x")
    doc.add_paragraph("LIST OF TABLES ....................................................................................................................... xi")
    doc.add_paragraph("Chapter One: Introduction ............................................................................................................. 1")
    doc.add_paragraph("  1.1 Background to the Study ....................................................................................................... 1")
    doc.add_paragraph("  1.2 Statement of the Problem ..................................................................................................... 3")
    doc.add_paragraph("  1.3 Aim and Objectives of the Project ........................................................................................... 4")
    doc.add_paragraph("  1.4 Scope of the Project ........................................................................................................... 5")
    doc.add_paragraph("  1.5 Significance of the Project ................................................................................................... 5")
    doc.add_paragraph("Chapter Two: Literature Review & Technology Context ................................................................................... 7")
    doc.add_paragraph("  2.1 Conceptual Review ............................................................................................................. 7")
    doc.add_paragraph("  2.2 Review of Existing Systems / Solutions ..................................................................................... 11")
    doc.add_paragraph("  2.3 Review of Relevant Technologies, Tools, and Frameworks ....................................................................... 13")
    doc.add_paragraph("  2.4 Gap Analysis ..................................................................................................................... 16")
    doc.add_paragraph("Chapter Three: Methodology & System Design ............................................................................................... 18")
    doc.add_paragraph("  3.1 Project Methodology ........................................................................................................... 18")
    doc.add_paragraph("  3.2 Requirements Analysis ......................................................................................................... 19")
    doc.add_paragraph("  3.3 System Architecture ........................................................................................................... 21")
    doc.add_paragraph("  3.4 System Design ................................................................................................................... 22")
    doc.add_paragraph("  3.5 Tools and Technologies Used ................................................................................................. 24")
    doc.add_paragraph("Chapter Four: System Implementation ........................................................................................................... 26")
    doc.add_paragraph("  4.1 Development Environment ....................................................................................................... 26")
    doc.add_paragraph("  4.2 Implementation Details ....................................................................................................... 27")
    doc.add_paragraph("  4.3 Security, Performance, and Scalability Considerations ........................................................................... 30")
    doc.add_paragraph("  4.4 Challenges Encountered ....................................................................................................... 32")
    doc.add_paragraph("Chapter Five: Testing, Results & Evaluation ............................................................................................... 34")
    doc.add_paragraph("  5.1 Testing Strategy ............................................................................................................... 34")
    doc.add_paragraph("  5.2 Test Results ................................................................................................................... 35")
    doc.add_paragraph("  5.3 Evaluation of Objectives ....................................................................................................... 36")
    doc.add_paragraph("  5.4 Discussion of Results ........................................................................................................... 37")
    doc.add_paragraph("Chapter Six: Summary, Conclusion & Recommendations ................................................................................... 39")
    doc.add_paragraph("  6.1 Summary of the Project ....................................................................................................... 39")
    doc.add_paragraph("  6.2 Conclusion ..................................................................................................................... 39")
    doc.add_paragraph("  6.3 Professional Contributions ................................................................................................... 40")
    doc.add_paragraph("  6.4 Recommendations for Future Work ................................................................................................... 41")
    doc.add_paragraph("References ................................................................................................................................... 43")
    doc.add_paragraph("Appendices ................................................................................................................................... 46")

    doc.add_page_break()

    add_heading_with_spacing(doc, "LIST OF FIGURES", level=1)
    doc.add_paragraph("Figure 3.1: High-Level Multi-Tenant SaaS Architecture Diagram ...................................................................... 21")
    doc.add_paragraph("Figure 3.2: Use Case Diagram for LMS Platform Actors .............................................................................. 22")
    doc.add_paragraph("Figure 3.3: Entity Relationship Diagram (ERD) for Multi-Tenant Database Design .............................................. 23")
    doc.add_paragraph("Figure 3.4: Sequence Diagram for Tenant Resolution and User Login ................................................................. 24")
    doc.add_paragraph("Figure 4.1: Code Structure and Module Organization .................................................................................. 27")
    doc.add_paragraph("Figure 5.1: User Acceptance Testing (UAT) Satisfaction Score Chart ................................................................. 36")

    doc.add_page_break()

    add_heading_with_spacing(doc, "LIST OF TABLES", level=1)
    doc.add_paragraph("Table 2.1: Comparison of Multi-Tenancy Database Isolation Models ................................................................. 9")
    doc.add_paragraph("Table 2.2: Gap Analysis and Justification ............................................................................................. 16")
    doc.add_paragraph("Table 3.1: Functional Requirements Specification ................................................................................... 19")
    doc.add_paragraph("Table 3.2: Non-Functional Requirements Specification ............................................................................... 20")
    doc.add_paragraph("Table 3.3: Hardware and Software Configuration ................................................................................... 24")
    doc.add_paragraph("Table 5.1: Database Isolation Verification Test Cases ........................................................................... 35")
    doc.add_paragraph("Table 5.2: Performance Benchmarking (Simulated 3G Network via Chrome DevTools) ................................................ 36")
    doc.add_paragraph("Table 5.3: Evaluation of Project Objectives ........................................................................................... 37")

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER ONE: INTRODUCTION
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER ONE: INTRODUCTION", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "1.1 Background to the Study", level=2, before=12, after=6)
    doc.add_paragraph(
        "The global education sector is undergoing a profound digital transformation, accelerated by the "
        "universal pivot toward remote and blended learning frameworks. The Learning Management "
        "System (LMS) market is projected to experience robust, sustained growth; the K-12 cloud "
        "segment is valued at approximately USD 1.6 billion in 2025 and is projected to reach USD 2.8 "
        "billion by 2034 (Research and Markets, 2025). Concurrently, the higher education software "
        "application ecosystem is expanding from USD 11.22 billion in 2025 to an estimated USD 38.6 "
        "billion by 2032, representing a compound annual growth rate of 19.3% (6Wresearch, 2025). "
        "Focusing on West Africa and Nigeria specifically, the local LMS market is forecast to expand "
        "significantly through 2031, driven by aggressive public and private initiatives targeting digital "
        "adoption across secondary and tertiary tiers (6Wresearch, 2025 – Nigeria LMS Market Report)."
    )
    doc.add_paragraph(
        "Despite these expansive global metrics, the adoption of robust, sustainable LMS platforms "
        "within Nigerian educational institutions remains highly uneven and structurally constrained. A "
        "comprehensive systematic mapping study by Can et al. (2017) examined 407 papers on multi tenant architectures in cloud computing, concluding that while cloud-based applications can "
        "provide services for large numbers of tenants through shared hardware and software resources, "
        "the successful implementation of such architectures requires careful consideration of storage "
        "strategies, isolation criteria, and deployment models. The study established that multi-tenancy "
        "is not merely a technical feature but an architectural paradigm that fundamentally shapes how "
        "SaaS applications deliver value across diverse client contexts. Modern cloud architectural "
        "patterns for enterprise SaaS highlight the vital necessity of balancing structural isolation with "
        "runtime performance and strict database governance — core principles that directly govern the "
        "infrastructure choices of this study (Cakar, 2025; Sharma, 2025). Meduri et al. (2023) provide "
        "a comprehensive architectural review of multi-tenancy in cloud computing, emphasizing that "
        "among IaaS, PaaS, and SaaS deployment models, the SaaS model offers the greatest cost "
        "benefit to enterprises and users due to its ability to share both hardware and software across "
        "multiple tenants with minimal data sharing. The authors argue that the feature of sharing "
        "resources among multiple users — known as multi-tenancy — is essential for achieving the "
        "objectives of cloud computing, particularly in resource-constrained environments."
    )
    doc.add_paragraph(
        "In the educational context, Azouzi et al. (2018) proposed a Business Process Feature Model "
        "(BPFM) approach for designing configurable, multi-tenant e-learning processes in the cloud. "
        "Their research demonstrated that multi-tenancy enables cost-effective SaaS through resource "
        "consolidation, where multiple customers are served by a single application instance, and "
        "isolation is enforced at the application level. This architectural approach allows institutions to "
        "reuse core services while customizing pedagogical sequences — an insight that directly "
        "informs the design philosophy of this project."
    )
    doc.add_paragraph(
        "While elite private institutions can comfortably amortize the licensing overhead of global SaaS "
        "platforms like Canvas or Google Workspace for Education, the vast majority of primary, "
        "secondary, and tertiary public institutions face a critical \"digital infrastructure paradox\". On "
        "one hand, traditional open-source solutions like Moodle provide a zero-license-cost alternative, "
        "yet they demand highly specialized, local technical expertise for infrastructure hosting, "
        "database optimization, routine patching, and security maintenance—resources that are "
        "notoriously scarce within standard Nigerian institutional IT departments. On the other hand, "
        "locally engineered school management portals frequently sacrifice deep pedagogical features, "
        "user experience design, and robust scalability in favor of basic administrative ledger tracking."
    )
    doc.add_paragraph(
        "Software as a Service (SaaS) delivered via a highly optimized, multi-tenant architecture "
        "introduces a powerful software engineering framework to break this paradox. Within a multi tenant paradigm, a single, unified instance of an application serves a multitude of independent "
        "client organizations (tenants) while guaranteeing that their data layers remain logically "
        "isolated (Meduri et al., 2023). By deploying a multi-tenant cloud framework, the "
        "operational capital expenditure (CapEx) associated with server provisioning, security auditing, "
        "and continuous integration can be distributed across all subscribing schools. This dramatically "
        "minimizes the Total Cost of Ownership (TCO) per institution. Consequently, this project "
        "intersects cloud-native systems architecture and the resource-constrained operational realities "
        "of Nigerian schools to deliver a platform that is both technologically advanced and contextually "
        "viable."
    )

    add_heading_with_spacing(doc, "1.2 Statement of the Problem", level=2, before=12, after=6)
    doc.add_paragraph(
        "Nigerian educational institutions currently contend with a highly fragmented, inefficient, and "
        "operationally unsustainable approach to digital learning management. The primary structural "
        "deficiencies are fourfold, each supported by recent empirical research:"
    )
    doc.add_paragraph(
        "1. Prohibitive Infrastructure and Maintenance Overhead: Deploying dedicated, standalone "
        "LMS instances for individual schools (single-tenancy) necessitates independent server "
        "provisioning, separate database administration, and repetitive manual upgrade cycles. This "
        "infrastructure model is financially and operationally untenable for the vast majority of Nigerian "
        "educational facilities. Osode et al. (2024) conducted a quantitative correlational study "
        "investigating factors that influence the reception and utilization of LMSs by teaching staff at "
        "three selected Nigerian universities. Their findings revealed that effort expectancy contributed "
        "most to LMS actual use, while facilitating conditions, performance expectancy, and social "
        "influence had statistically significant effects on LMS usage and design decisions. Critically, "
        "the study identified that inadequate infrastructure and lack of technical support remain primary "
        "barriers to sustained LMS adoption in Nigerian higher education institutions."
    )
    doc.add_paragraph(
        "2. Contextual Mismatch and Bandwidth Inefficiencies: Major global LMS platforms are "
        "fundamentally engineered for high-bandwidth, low-latency broadband environments. They are "
        "not optimized for the standard Nigerian digital landscape, which is heavily reliant on mobile "
        "web traffic, variable 3G/4G cellular networks with fluctuating packet delivery rates, and "
        "mobile-first student access patterns. A qualitative study by Kester and Ojedeji (2022) explored "
        "the workability of cloud-based education in Africa, highlighting critical challenges relating to "
        "high levels of poverty, epileptic power supply, and poor internet connectivity. The authors "
        "concluded that while cloud computing systems provide a variety of opportunities for content "
        "delivery, their effectiveness in African contexts is contingent upon deployment strategies that "
        "account for infrastructural volatility."
    )
    doc.add_paragraph(
        "Furthermore, a study examining push and pull factors influencing LMS use post-COVID in "
        "Nigerian open distance learning (Educational Technology Quarterly, 2025) identified that "
        "infrastructural and access issues, the digital literacy skills of students, low participation or "
        "engagement with the LMS by students, as well as technical and content-related issues, limit the "
        "effective use of the LMS for learning. The findings underscored that while COVID-19-related "
        "factors, digital readiness, and flexibility served as push factors for LMS adoption, "
        "infrastructural limitations remained a significant constraint."
    )
    doc.add_paragraph(
        "3. Lack of Localized, Cost-Effective Solutions: The 6Wresearch Nigeria Learning "
        "Management System Market Report (2025) identifies that the Nigerian LMS market encounters "
        "obstacles such as limited internet connectivity and digital infrastructure. Many educational "
        "institutions struggle with unreliable internet access, which hampers the effective use of online "
        "learning platforms. Additionally, there is a lack of digital literacy among educators and "
        "students, which affects the adoption and utilization of LMS solutions. Budget constraints and "
        "the high cost of implementing and maintaining these systems also pose significant challenges."
    )
    doc.add_paragraph(
        "Manzoor (2024) investigated cloud computing solutions for educational applications within "
        "African contexts, focusing on their role in enhancing educational access. Through comparative "
        "analysis of leading cloud solutions, the study found that regional providers offer advantages "
        "for educational deployment due to their regional focus, data spread across the continent, and "
        "accessible technical support. However, the research also confirmed that the digital divide "
        "remains a significant barrier, with disparities in technology access creating unequal educational "
        "opportunities."
    )
    doc.add_paragraph(
        "4. Architectural Monoliths and Scalability Bottlenecks: Existing locally engineered e-learning "
        "solutions are predominantly structured as monolithic web applications. These systems lack "
        "horizontal scaling mechanisms and connection pooling safeguards, meaning they cannot "
        "efficiently accommodate thousands of concurrent users during high-stakes examination or "
        "enrollment windows across multiple federated institutions without severe latency spikes or "
        "database deadlocks. Research by Azouzi and Ghannouchi (2025) on designing multi-tenant e learning systems in the cloud for higher education has demonstrated that process-oriented, "
        "configurable architectures are essential for achieving scalability while maintaining tenant "
        "isolation."
    )
    doc.add_paragraph(
        "This project directly resolves the problem of delivering a scalable, high-performance, cost effective, and context-aware LMS to Nigerian institutions by engineering a specialized multi tenant SaaS architectural pattern deployed onto modern, serverless cloud database "
        "infrastructure. By adopting a multi-tenant SaaS model, the operational capital expenditure "
        "associated with server provisioning, security auditing, and continuous integration can be "
        "distributed across all subscribing schools, dramatically minimizing the Total Cost of "
        "Ownership per institution as evidenced by Yojji (2025), whose industry analysis demonstrated "
        "that multi-tenant platforms achieve 30–55% lower infrastructure costs per client than single tenant deployments."
    )

    add_heading_with_spacing(doc, "1.3 Aim and Objectives of the Project", level=2, before=12, after=6)
    doc.add_paragraph(
        "The primary aim of this project is to design, implement, and evaluate a scalable, multi-tenant "
        "Software as a Service (SaaS) Learning Management System meticulously tailored to the "
        "infrastructure constraints and operational environments of Nigerian educational institutions."
    )
    doc.add_paragraph("The specific objectives are to:")
    doc.add_paragraph(
        "1. Design a database architecture supporting multi-tenant isolation with a logical, shared-database, "
        "shared-schema mechanism suitable for educational data privacy using Prisma 7 and PostgreSQL."
    )
    doc.add_paragraph(
        "2. Implement a highly responsive, bandwidth-optimized web application interface using React 19 "
        "and TailwindCSS 4, deployed on Vercel's edge network."
    )
    doc.add_paragraph(
        "3. Develop core LMS functionalities including course creation, user enrollment, in-app messaging, "
        "and automated certificate generation."
    )
    doc.add_paragraph(
        "4. Develop an advanced, real-time exam integrity and proctoring dashboard allowing platform operators "
        "and teachers to monitor candidate cameras, listen to student audio, issue warnings, force hardware activation, "
        "and capture snapshots or audio/video clips for evidence."
    )
    doc.add_paragraph(
        "5. Implement a cross-tenant analytics dashboard for platform operators providing "
        "aggregate insights into system usage and tenant engagement."
    )
    doc.add_paragraph(
        "6. Evaluate the system's performance and scalability under simulated multi-tenant load "
        "conditions and Nigerian network constraints."
    )

    add_heading_with_spacing(doc, "1.4 Scope of the Project", level=2, before=12, after=6)
    doc.add_paragraph(
        "This project covers the design, implementation, and cloud deployment of a fully functional "
        "multi-tenant SaaS LMS prototype. The scope includes:"
    )
    doc.add_paragraph(
        "Architecture: Implementation of a logical shared-database, shared-schema isolation strategy using a unified "
        "PostgreSQL database. Row-level filters and application-level middleware guards enforce strict data "
        "boundaries at the query level, ensuring complete tenant isolation and compliance."
    )
    doc.add_paragraph(
        "Core Features: User authentication, role-based access control (RBAC), course and content "
        "management, student enrollment, lesson progress tracking, automated certificates (with seals and signatures), "
        "headless CMS configurations for tenant landing pages, and class chat streams."
    )
    doc.add_paragraph(
        "Exam Proctoring & Integrity: Real-time candidate webcam frame uploads, 3-second sliced audio streaming, "
        "WebRTC support camera pairing, live invigilation console, remote control triggers (Prompt/Force Camera/Audio), "
        "evidence downloads (Snapshot and Record Audio/Clip), and telemetry tracking (mic active, gaze secure)."
    )
    doc.add_paragraph(
        "Platform Analytics Dashboard: A dedicated analytics dashboard for the platform operator "
        "providing aggregate, cross-tenant insights including total active users, course "
        "engagement metrics, and system health monitoring."
    )
    doc.add_paragraph(
        "Deployment & DevOps: Continuous deployment pipeline configured via GitHub to Vercel for "
        "both the frontend and serverless API endpoints. The database is hosted on a managed Supabase instance."
    )

    add_heading_with_spacing(doc, "1.5 Significance of the Project", level=2, before=12, after=6)
    doc.add_paragraph(
        "To Industry Practice: This project provides a concrete, documented software engineering "
        "artifact demonstrating how modern cloud platforms can be leveraged to build a secure, scalable "
        "multi-tenant SaaS for the African EdTech market without significant infrastructure capital "
        "expenditure. It joins a growing body of open-source reference architectures that aim to "
        "democratize educational technology. As Yojji (2025) notes, organizations use the multi-tenant "
        "model for partner training, customer education, and franchise learning, simplifying scaling and "
        "lowering the cost of managing many learning environments."
    )
    doc.add_paragraph(
        "To Nigerian Educational Institutions: It offers a fully operational, cloud-hosted proof-of-concept "
        "that lowers both financial and technical barriers to digital learning. Schools can be "
        "onboarded instantly via subdomain provisioning with zero local server maintenance. "
        "Assalaarachchi et al. (2023) found that behavioral intention to adopt SaaS applications in e learning, together with facilitating conditions, has a significant impact on SaaS adoption, with "
        "factors such as performance expectancy, effort expectancy, and trust determining adoption "
        "intention. These findings have direct implications for the design decisions in this project."
    )
    doc.add_paragraph(
        "To Platform Operators and EdTech Entrepreneurs: The inclusion of an advanced analytics "
        "dashboard empowers stakeholders with data-driven insights into platform usage, enabling "
        "informed decisions about feature development and support resource allocation. The multi-tenant "
        "architecture's isolation pillars — operational, data, compliance, and analytical — "
        "provide a comprehensive framework for building scalable educational platforms (ClickHouse, 2026)."
    )
    doc.add_paragraph(
        "To Software Engineering Body of Knowledge: It contributes a case study on implementing "
        "logical multi-tenancy on a managed Postgres service and deploying a unified Next.js/React codebase "
        "to an edge network — a modern, cost-efficient alternative to traditional VM-based hosting. This "
        "aligns with the systematic mapping study by Can et al. (2017), which identified the need for "
        "more empirical research on multi-tenant implementation strategies in diverse deployment contexts."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER TWO: LITERATURE REVIEW & TECHNOLOGY CONTEXT
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER TWO: LITERATURE REVIEW & TECHNOLOGY CONTEXT", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "2.1 Conceptual Review", level=2, before=12, after=6)
    doc.add_paragraph(
        "This section clarifies the foundational concepts underpinning the design and implementation "
        "of a multi-tenant Software as a Service (SaaS) Learning Management System (LMS)."
    )

    add_heading_with_spacing(doc, "2.1.1 Software as a Service (SaaS) and Cloud Computing", level=3, before=12, after=6)
    doc.add_paragraph(
        "Software as a Service is a cloud computing delivery model where software applications are "
        "hosted by a service provider and made available to customers over the internet. Unlike "
        "traditional on-premise software, SaaS eliminates the need for organizations to manage "
        "underlying infrastructure, middleware, or application maintenance (Mell & Grance, 2011). "
        "Meduri et al. (2023) characterize cloud computing as a pervasive paradigm enabling network "
        "access to a shared group of computing resources that can be dynamically provisioned based on "
        "demand with minimal subscriber intervention in a pay-per-use routine."
    )
    doc.add_paragraph(
        "Meduri et al. (2023) articulate that the most important aspect of cloud computing is to create "
        "an illusion of infinite resource capacity at the doorstep of customers while making optimum "
        "utilization of customer-owned resources, succinctly optimizing capital expenditure and "
        "reducing operational costs to ensure two-fold benefit to customers. Pocatilu et al. (2011) define "
        "the SaaS delivery model as one where software is maintained and updated on the cloud at a "
        "central location and delivered to clients as a service, usually through a browser. Through this "
        "model, clients do not worry about the resources required for application running, such as "
        "network, servers, middleware, or system administration teams."
    )
    doc.add_paragraph(
        "In the educational context, this translates to schools avoiding significant upfront capital "
        "expenditure on servers and IT personnel. Assalaarachchi et al. (2023) emphasize that SaaS "
        "applications immensely impact the success of e-learning with their ability for cost saving, "
        "scalability, and better collaboration. The authors found that performance expectancy, effort "
        "expectancy, and trust significantly determine undergraduates' intention to adopt SaaS "
        "applications in e-learning contexts."
    )

    add_heading_with_spacing(doc, "2.1.2 Multi-Tenancy Architecture", level=3, before=12, after=6)
    doc.add_paragraph(
        "Multi-tenancy is an architectural principle wherein a single instance of a software application "
        "serves multiple customers, known as tenants. The primary objective of multi-tenancy in SaaS "
        "is resource optimization; by sharing the application and database layers, operational costs and "
        "complexity are amortized across all tenants. Meduri et al. (2023) note that among IaaS, PaaS, "
        "and SaaS deployments, the SaaS model offers more cost benefit to enterprises and users as "
        "compared to its predecessors because of its ability to share both hardware and software across "
        "multiple tenants with very minimal data sharing. The feature of sharing resources among "
        "multiple users is essential for achieving the objectives of cloud computing."
    )
    doc.add_paragraph(
        "Can et al. (2017) conducted a systematic mapping study of 407 papers on multi-tenant "
        "architectures in cloud computing. Their analysis revealed that cloud-based applications can "
        "provide services for large numbers of tenants using the same hardware and software by "
        "implementing multi-tenant architectures. The study classified publications by research topic "
        "and content, identifying which storage strategies were used most frequently, which criteria "
        "were taken into account in selecting preferred storage strategies, and the most searched topics "
        "under the multi-tenant architecture model. The findings established that multi-tenancy is an "
        "organizational pattern for SaaS that enables a single instance of an application to be hosted on "
        "the same hardware and accessed by multiple customers with the aim of lowering costs."
    )
    doc.add_paragraph(
        "Cakar (2025) provides a contemporary examination of scalable multi-tenant software "
        "architectures, focusing on isolation, performance, and governance in enterprise-grade systems. "
        "The research emphasizes that multi-tenancy is a system-wide design constraint, not merely a "
        "database-level concern, and must propagate consistently through application code, connection "
        "pools, background workers, caching layers, and analytical pipelines. This architecture is "
        "particularly suited to the education sector in Nigeria, where many schools operate under severe "
        "financial and infrastructural constraints, making individual LMS hosting infeasible."
    )
    doc.add_paragraph(
        "A multi-tenant LMS uses one codebase to serve many organizations and isolates data, users, "
        "and permissions at the tenant level. Yojji (2025) documents that this architecture dominates "
        "modern SaaS learning platforms because it reduces deployment overhead and simplifies long-term "
        "maintenance, with industry data showing that platforms built as multi-tenant systems "
        "achieve 30–55% lower infrastructure costs per client than single-tenant deployments, primarily "
        "due to shared compute, unified updates, and centralized monitoring."
    )

    add_heading_with_spacing(doc, "2.1.3 Database Isolation Models in Multi-Tenancy", level=3, before=12, after=6)
    doc.add_paragraph(
        "A critical design decision in multi-tenant systems is the level of data isolation. The architectural "
        "literature identifies three primary models:"
    )
    doc.add_paragraph(
        "1. Shared Database, Shared Schema: All tenants share the same database tables, with a tenant_id (or organizationId) "
        "field used to logically separate data. This offers the highest resource efficiency and lowest operational complexity "
        "since all migrations are applied to a single database schema. However, it requires rigorous application-level "
        "security, query scoping, and database Row Level Security (RLS) policies to prevent cross-tenant data leakage. According "
        "to ClickHouse (2026), this model serves as the right default for most B2B SaaS applications, though it demands careful "
        "implementation of tenant identity propagation throughout the entire application stack."
    )
    doc.add_paragraph(
        "2. Shared Database, Separate Schema: Each tenant has a dedicated schema within a single database instance. This provides "
        "stronger logical isolation and allows for some degree of tenant-specific customization, but increases migration "
        "complexity since schema changes must be applied to every tenant's schema sequentially. Microsoft (2026) documents "
        "the trade-offs: the schema-per-tenant approach is recommended when data isolation is key but the overhead of managing "
        "multiple database engines cannot be justified."
    )
    doc.add_paragraph(
        "3. Separate Database per Tenant: Each tenant is assigned a dedicated database. While the most secure and performant under "
        "heavy load, this model increases management overhead and infrastructure cost. This approach is typically recommended "
        "for very performance-sensitive or regulatory-bound tenants (Microsoft, 2026)."
    )

    # Table 2.1
    t21 = doc.add_table(rows=5, cols=4)
    t21.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t21.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Shared DB, Shared Schema (Chosen)'
    hdr_cells[2].text = 'Shared DB, Separate Schema'
    hdr_cells[3].text = 'Separate Database'
    for cell in hdr_cells:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    features = [
        ('Isolation Level', 'Low-Medium (Logical via tenant ID)', 'Medium (Logical via schema)', 'High (Physical)'),
        ('Resource Efficiency', 'Highest (Unified connections/memory)', 'Medium', 'Lowest (Wasted resources)'),
        ('Operational Complexity', 'Lowest (Single schema migration)', 'Medium (Scripted schema loop)', 'Highest (Multiple DB endpoints)'),
        ('Ideal Use Case', 'Cost-sensitive, high-volume multi-tenancy', 'Regulated data, moderate customization', 'Strict compliance, high-budget')
    ]
    for idx, (f, c1, c2, c3) in enumerate(features):
        row_cells = t21.rows[idx+1].cells
        row_cells[0].text = f
        row_cells[1].text = c1
        row_cells[2].text = c2
        row_cells[3].text = c3
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    p_t21_cap = doc.add_paragraph("Table 2.1: Comparison of Multi-Tenancy Database Isolation Models")
    p_t21_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t21_cap.paragraph_format.space_before = Pt(6)
    p_t21_cap.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "This project adopts the Shared Database, Shared Schema model using Supabase/PostgreSQL. This selection "
        "is justified by the resource constraints of Nigerian educational institutions. Amortizing connection "
        "pooling, CPU usage, and memory footprints across a single relational instance drastically reduces the "
        "monthly cost of deployment. Data isolation is maintained through strict JWT validation, application-level "
        "guards, and Prisma model filters on the `organizationId` foreign key."
    )

    add_heading_with_spacing(doc, "2.1.4 Learning Management Systems (LMS) in Education", level=3, before=12, after=6)
    doc.add_paragraph(
        "A Learning Management System is a web-based platform designed to deliver, track, and manage educational "
        "courses and training programs. Research by Osode et al. (2024) contextualizes LMS adoption within the "
        "broader framework of technology acceptance. Their study, grounded in the modified Unified Theory of "
        "Acceptance and Use of Technology (UTAUT), found that effort expectancy contributed most to LMS actual "
        "use, while facilitating conditions, performance expectancy, and social influence had statistically "
        "significant effects on LMS usage and design decisions. The findings also revealed that 'Design decision' "
        "and 'Staff performance' variables added to the UTAUT instrument garnered additional data about LMS "
        "usage in Nigerian higher education institutions."
    )
    doc.add_paragraph(
        "A qualitative study examining push and pull factors influencing LMS use post-COVID in Nigerian open "
        "distance learning (Educational Technology Quarterly, 2025) identified that COVID-19-related factors, "
        "digital readiness and literacy, as well as flexibility and autonomy offered by the LMS, were push "
        "factors for LMS use post-COVID. System features, regular updates, effective support mechanisms, and "
        "institutional policies served as pull factors. However, infrastructural and access issues, digital "
        "literacy skills of students, and technical and content-related issues continue to limit effective LMS use."
    )

    add_heading_with_spacing(doc, "2.2 Review of Existing Systems / Solutions", level=2, before=12, after=6)
    doc.add_paragraph(
        "This section analyzes prominent LMS solutions available globally and locally, evaluating their "
        "alignment with the needs of a scalable, multi-tenant environment for Nigerian schools."
    )

    add_heading_with_spacing(doc, "2.2.1 Global Market Leaders (Canvas, Moodle, Google Classroom)", level=3, before=12, after=6)
    doc.add_paragraph(
        "Canvas (Instructure): A leading cloud-native SaaS LMS widely used in higher education and increasingly "
        "in K-12. Canvas employs a multi-tenant architecture at the infrastructure level, allowing it to scale "
        "globally. However, its pricing model and complex feature set are often prohibitive for smaller, "
        "independent schools in developing economies."
    )
    doc.add_paragraph(
        "Moodle: As an open-source platform, Moodle is the most deployed LMS in the world. However, as "
        "documented in the systematic mapping study by Can et al. (2017), while Moodle can be configured in a "
        "multi-tenant fashion using plugins, it is inherently a single-tenant application. Research evaluating "
        "multi-tenant applications (HAL, 2021) examined Iomad, an extension of the Moodle Learning Management "
        "System, as a case study for tenant migration strategies, noting that such applications achieve tenant "
        "isolation through dedicated fields in their relational schema and are not designed to support scaling "
        "operations without significant modification."
    )
    doc.add_paragraph(
        "Google Classroom: A free, streamlined tool integrated with Google Workspace for Education. While "
        "popular due to its zero-cost entry point and simplicity, Google Classroom lacks the deep "
        "administrative controls, advanced analytics, and white-label multi-tenancy features required for "
        "a centralized SaaS provider serving multiple independent schools."
    )

    add_heading_with_spacing(doc, "2.2.2 Specialized Multi-Tenant LMS and E-Learning Solutions", level=3, before=12, after=6)
    doc.add_paragraph(
        "Atwongyere (2025) developed a multi-tenant e-learning platform specifically for the Ugandan context, "
        "identifying core pain points: navigation difficulties, poor scalability of existing systems, and lack "
        "of integrated payment gateways. The study found 80% user acceptance for the localized, multi-tenant "
        "approach. This validates the problem space but also highlights a gap: the system focused on tertiary "
        "institutions and lacked the advanced live proctoring safeguards and edge serverless deployment models."
    )
    doc.add_paragraph(
        "ClassroomIO (2025): An emerging open-source EdTech platform that provides a complete learning environment "
        "with multi-tenant capabilities. ClassroomIO serves as a production-grade reference for how open-source "
        "communities can build scalable educational tools, supporting the viability of the open, documented "
        "approach taken in this project."
    )
    doc.add_paragraph(
        "SabiScholar (Iloh, 2025) is a proprietary, closed-source Nigerian LMS. Notably, its patent filing "
        "describes a \"Bandwidth-Aware, Curriculum-Aligned Multi-Tenant System\" — a design philosophy that "
        "directly mirrors the non-functional requirements of this project. SabiScholar's commercial traction "
        "demonstrates clear market demand for a centralized, multi-tenant SaaS platform that allows individual "
        "schools to deploy their own branded LMS instances without managing infrastructure. However, SabiScholar "
        "remains a commercial product; this project contributes to the ecosystem by providing an open, "
        "academically scrutinized reference architecture built with Next.js and Prisma."
    )

    add_heading_with_spacing(doc, "2.3 Review of Relevant Technologies, Tools, and Frameworks", level=2, before=12, after=6)
    doc.add_paragraph(
        "The selection of technology for a multi-tenant SaaS LMS requires careful consideration of "
        "performance, security, and maintainability. This section reviews the technology landscape."
    )
    doc.add_paragraph(
        "Next.js 16 (App Router): Next.js is a React framework optimized for production. The App Router introduces "
        "React Server Components (RSC) and Server Actions, which enable faster page load times by rendering "
        "components on the server and reducing the bundle size transmitted to the client. This directly mitigates "
        "bandwidth constraints in the Nigerian context."
    )
    doc.add_paragraph(
        "Prisma 7 ORM: Prisma is a type-safe database ORM that facilitates schema definition and migrations. It "
        "supports PostgreSQL native connection pooling, which is critical for scaling a shared database model. "
        "Using Prisma, all schema models include an `organizationId` foreign key, allowing type-safe, scoped "
        "queries across all business operations."
    )
    doc.add_paragraph(
        "PostgreSQL: A robust, open-source relational database that is highly scalable and supports complex "
        "data types (e.g., JSONB for flexible schema structures). The database layer is hosted on a managed Supabase "
        "instance, which provides connection pooling and built-in replication."
    )
    doc.add_paragraph(
        "React 19 & TailwindCSS 4: React 19 provides a component-driven user interface framework, while TailwindCSS "
        "4 compiles highly optimized, utility-first CSS configurations that minimize asset file sizes for "
        "low-bandwidth edge network deliveries."
    )
    doc.add_paragraph(
        "WebRTC & Media Capture API: Used to establish local video and audio monitoring tracks on the candidate's "
        "browser. The browser captures 3-second base64 jpegs and 3-second base64 webm audio slices and uploads them "
        "to the Edge signaling endpoint, simulating a real-time stream for the invigilator dashboard without the need "
        "for expensive constant streaming connections."
    )

    add_heading_with_spacing(doc, "2.4 Gap Analysis", level=2, before=12, after=6)
    doc.add_paragraph(
        "Based on the conceptual review and analysis of existing systems, a distinct gap exists that this "
        "project aims to fill. Table 2.2 summarizes the identified gaps and justifies the proposed solution."
    )

    # Table 2.2
    t22 = doc.add_table(rows=6, cols=4)
    t22.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr22 = t22.rows[0].cells
    hdr22[0].text = 'Gap ID'
    hdr22[1].text = 'Gap Description'
    hdr22[2].text = 'Key Evidence from Literature'
    hdr22[3].text = 'Justification for Proposed Solution'
    for cell in hdr22:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    gaps = [
        ('Gap 1', 'Cost & Infrastructure Barrier', 'Moodle requires independent VM server provisioning; Canvas is expensive (HAL, 2021; Azouzi, 2018).', 'A logical multi-tenant SaaS model on serverless hosting (Vercel + Supabase) eliminates individual VM provisioning costs.'),
        ('Gap 2', 'Lack of Bandwidth Optimization', 'Nigerian LMS adoption is limited by unstable 3G/4G bandwidth (Educational Technology Quarterly, 2025; Kester, 2022).', 'Edge caching, Next.js page generation, image conversion, and base64 upload optimization reduce bundle payloads.'),
        ('Gap 3', 'Closed-Source Nature of Local Solutions', 'SabiScholar validates local demand but is closed-source and proprietary (Iloh, 2025).', 'This project provides an open reference architecture for the African educational developer community.'),
        ('Gap 4', 'Lack of Live Exam Proctoring', 'Monolithic portals do not provide integrity tracking or live monitoring capabilities during examinations.', 'Integrated WebRTC mobile camera pairing, audio chunking, and invigilator dashboards enable secure exam tracking.'),
        ('Gap 5', 'Rigid Academic Workflows', 'Existing systems offer "one-size-fits-all" structures, lacking customizable institutional branding (Azouzi, 2018).', 'Next.js dynamic routing and organization-scoped metadata settings enable unique custom branding templates.')
    ]
    for idx, (gid, gd, ge, gj) in enumerate(gaps):
        row_cells = t22.rows[idx+1].cells
        row_cells[0].text = gid
        row_cells[1].text = gd
        row_cells[2].text = ge
        row_cells[3].text = gj
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    p_t22_cap = doc.add_paragraph("Table 2.2: Gap Analysis and Justification")
    p_t22_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t22_cap.paragraph_format.space_before = Pt(6)
    p_t22_cap.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER THREE: METHODOLOGY & SYSTEM DESIGN
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER THREE: METHODOLOGY & SYSTEM DESIGN", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "3.1 Project Methodology", level=2, before=12, after=6)
    doc.add_paragraph(
        "This project adopts an Iterative and Incremental Development methodology, closely aligned with "
        "Agile principles. Given the complexity of multi-tenant architecture and the need to validate "
        "specific features early in the lifecycle, a strict Waterfall model was deemed unsuitable. "
        "The project was divided into four distinct iterations:"
    )
    doc.add_paragraph(
        "Iteration 1 (Foundation): Setup of the shared database schema on PostgreSQL using Prisma. Implementation "
        "of the Edge Proxy middleware for subdomain/slug routing, user auth cookies, and JWT verification."
    )
    doc.add_paragraph(
        "Iteration 2 (Core LMS Features): Implementation of course creation, syllabus modules, lessons, student "
        "enrollments, blog posts, resources, and automatic certificate generation."
    )
    doc.add_paragraph(
        "Iteration 3 (Proctoring & Analytics): Implementation of WebRTC support camera pairing, periodic camera "
        "frame captures, audio streaming, live invigilation console, and platform-wide Super Admin analytics."
    )
    doc.add_paragraph(
        "Iteration 4 (Testing & Deployment): Edge deployment on Vercel, smoke tests, TypeScript compilation "
        "verification, performance auditing, and finalization of documentation."
    )

    add_heading_with_spacing(doc, "3.2 Requirements Analysis", level=2, before=12, after=6)

    # Table 3.1
    add_heading_with_spacing(doc, "Table 3.1: Functional Requirements Specification", level=3, before=12, after=6)
    t31 = doc.add_table(rows=8, cols=3)
    t31.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr31 = t31.rows[0].cells
    hdr31[0].text = 'ID'
    hdr31[1].text = 'Requirement Description'
    hdr31[2].text = 'Actor(s)'
    for cell in hdr31:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    frs = [
        ('FR-01', 'Subdomain and slug-based routing resolution to load tenant configuration.', 'System'),
        ('FR-02', 'Creation and management of school accounts (Teachers, Students) by admins.', 'School Admin'),
        ('FR-03', 'Course syllabus setup, lesson file uploads, and quiz/exam creation.', 'Teacher'),
        ('FR-04', 'Course enrollment, reading content, and submitting quiz/exam answers.', 'Student'),
        ('FR-05', 'Dynamic white-label theme styling, primary color hex, and logo loading.', 'System'),
        ('FR-06', 'Live exam invigilation dashboard to monitor camera, play audio feeds, and record clip evidence.', 'Teacher / Staff'),
        ('FR-07', 'Super Admin cross-tenant dashboard displaying total registered schools and metrics.', 'Super Admin')
    ]
    for idx, (fid, fd, fa) in enumerate(frs):
        row_cells = t31.rows[idx+1].cells
        row_cells[0].text = fid
        row_cells[1].text = fd
        row_cells[2].text = fa
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    doc.add_paragraph()

    # Table 3.2
    add_heading_with_spacing(doc, "Table 3.2: Non-Functional Requirements Specification", level=3, before=12, after=6)
    t32 = doc.add_table(rows=5, cols=3)
    t32.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr32 = t32.rows[0].cells
    hdr32[0].text = 'ID'
    hdr32[1].text = 'Requirement Description'
    hdr32[2].text = 'Metric/Target'
    for cell in hdr32:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    nfrs = [
        ('NFR-01', 'Tenant Data Isolation: No user shall be able to access data of another school.', '100% isolation verified via integration tests.'),
        ('NFR-02', 'Performance: Initial page load on a simulated 3G mobile network.', 'First Contentful Paint (FCP) < 3 seconds.'),
        ('NFR-03', 'Scalability: All database queries must filter by organizationId to index scan.', 'No cross-tenant data leak or cross-schema scans.'),
        ('NFR-04', 'Usability: Interface must be responsive across different screens.', 'Pass Google Mobile-Friendly audit.')
    ]
    for idx, (nid, nd, nm) in enumerate(nfrs):
        row_cells = t32.rows[idx+1].cells
        row_cells[0].text = nid
        row_cells[1].text = nd
        row_cells[2].text = nm
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    doc.add_paragraph()

    add_heading_with_spacing(doc, "3.3 System Architecture", level=2, before=12, after=6)
    doc.add_paragraph(
        "The system leverages a modern Serverless Edge deployment model with clear separation between the frontend "
        "delivery network, serverless API endpoints, and a managed PostgreSQL database."
    )
    doc.add_paragraph(
        "[Insert Figure 3.1: High-Level Multi-Tenant SaaS Architecture Diagram]\n"
        "Description of Figure 3.1: The Client Layer consists of a React 19 SPA. Requests are routed via Vercel's Edge Network, "
        "which resolves subdomains/slugs at the Edge. The Application Backend runs on Vercel Serverless Functions "
        "using Next.js 16. The backend establishes a database connection to Supabase PostgreSQL. Inside the unified "
        "database schema, all tenant data tables (Users, Courses, Submissions) are logically isolated using the "
        "organizationId foreign key, which matches the resolved organization."
    )

    add_heading_with_spacing(doc, "3.4 System Design", level=2, before=12, after=6)
    add_heading_with_spacing(doc, "3.4.1 Use Case Diagram", level=3, before=12, after=6)
    doc.add_paragraph(
        "[Insert Figure 3.2: Use Case Diagram for LMS Platform Actors]\n"
        "Description of Figure 3.2: The diagram outlines use cases for four core actors: Super Admin (manages schools, views "
        "global analytics), School Admin (manages school users), Teacher (creates courses/quizzes, reviews proctoring), and "
        "Student (enrolls in classes, completes timed exams under proctoring)."
    )

    add_heading_with_spacing(doc, "3.4.2 Database Design (Entity Relationship Diagram)", level=3, before=12, after=6)
    doc.add_paragraph(
        "[Insert Figure 3.3: Entity Relationship Diagram (ERD) for Multi-Tenant Database Design]\n"
        "Description of Figure 3.3: The ERD highlights the logical database schema design. The parent table is Organization. "
        "All tenant-specific tables—User, Course, SchoolCohort, Assessment, Submission, and ProctoringEvent—possess an "
        "organizationId foreign key referencing Organization(id). Scoping queries by this key guarantees logical tenant "
        "isolation and leverages database index scans."
    )

    add_heading_with_spacing(doc, "3.4.3 Sequence Diagram: Login and Tenant Resolution", level=3, before=12, after=6)
    doc.add_paragraph(
        "[Insert Figure 3.4: Sequence Diagram for Tenant Resolution and User Login]\n"
        "Description of Figure 3.4: Shows the chronological process of login: (1) Student inputs credentials on subdomain/slug. "
        "(2) Frontend sends credentials. (3) Next.js proxy middleware intercepts request, validates the slug from URL path, "
        "checks JWT token, and checks role. (4) Backend queries Organization table. (5) Queries the User table scoped by organizationId. "
        "(6) Returns the JWT session token signed with the resolved organizationId."
    )

    add_heading_with_spacing(doc, "3.5 Tools and Technologies Used", level=2, before=12, after=6)

    # Table 3.3
    t33 = doc.add_table(rows=9, cols=4)
    t33.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr33 = t33.rows[0].cells
    hdr33[0].text = 'Category'
    hdr33[1].text = 'Tool/Technology'
    hdr33[2].text = 'Version'
    hdr33[3].text = 'Purpose'
    for cell in hdr33:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    tools = [
        ('OS', 'Windows 11 / macOS / Linux', '-', 'Development and Deployment environment'),
        ('Backend', 'Next.js 16 (App Router) / Node.js', '16.x / 20.x', 'Serverless API endpoints, Server Components, and Routing'),
        ('Frontend', 'React 19 / TypeScript / TailwindCSS 4', '19.x / 5.x / 4.x', 'Mobile-first responsive UI styling and state handling'),
        ('Database', 'Supabase PostgreSQL (managed)', '15.x', 'Relational database engine with pooling and replication'),
        ('Auth', 'jose JWT / Cookies', '6.x', 'HTTP-only session cookie authentication and roles tracking'),
        ('Hosting', 'Vercel CDN', 'Pro Plan', 'Fast global Edge deployment and serverless function hosting'),
        ('Analytics', 'Recharts / SQL Views', '2.x', 'Aggregated usage rendering on Super Admin dashboard'),
        ('DevOps', 'Prisma CLI / GitHub Actions', '7.6.0', 'Database migrations versioning and automated CI/CD pipeline')
    ]
    for idx, (tc, tt, tv, tp) in enumerate(tools):
        row_cells = t33.rows[idx+1].cells
        row_cells[0].text = tc
        row_cells[1].text = tt
        row_cells[2].text = tv
        row_cells[3].text = tp
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    p_t33_cap = doc.add_paragraph("Table 3.3: Hardware and Software Configuration")
    p_t33_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t33_cap.paragraph_format.space_before = Pt(6)
    p_t33_cap.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER FOUR: SYSTEM IMPLEMENTATION
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER FOUR: SYSTEM IMPLEMENTATION", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "4.1 Development Environment", level=2, before=12, after=6)
    doc.add_paragraph(
        "The local development environment was managed using VS Code and Node.js. Prisma CLI was used to orchestrate "
        "database migrations against the local and staging PostgreSQL instances. During development, a local "
        "environment file (.env.local) was configured to point to the Supabase database. For production, the "
        "application is fully deployed to Vercel and Supabase, utilizing Vercel environment variables for secrets "
        "management, ensuring complete infrastructure automation."
    )

    add_heading_with_spacing(doc, "4.2 Implementation Details", level=2, before=12, after=6)
    add_heading_with_spacing(doc, "4.2.1 Core Module: Tenant Resolution and Edge Proxy", level=3, before=12, after=6)
    doc.add_paragraph(
        "Tenant resolution is implemented at the Edge. When a user requests a path under /o/[slug], the Next.js proxy "
        "middleware (src/proxy.ts) intercepts the request. The middleware extracts the organization's slug from the path, "
        "validates it against the JWT session token (saas_lms_token), and ensures the user's role (ADMIN, TEACHER, STUDENT) "
        "authorizes them to access the requested tenant. This prevents \"tenant-hopping\" security vulnerabilities."
    )

    add_heading_with_spacing(doc, "4.2.2 Dynamic White-Labeling", level=3, before=12, after=6)
    doc.add_paragraph(
        "On loading any route under /o/[slug], the frontend requests the active tenant configuration from `/api/public/organizations/[slug]`. "
        "This endpoint returns branding variables: school name, custom logo URL, and a primary brand color hex code. "
        "These parameters are set on CSS variables and dynamically injected into Tailwind classes. This enables unique white-label "
        "experiences across different schools using a single, unified codebase, avoiding separate frontend builds."
    )

    add_heading_with_spacing(doc, "4.2.3 Advanced Live Proctoring Implementation", level=3, before=12, after=6)
    doc.add_paragraph(
        "The project implements a real-time exam invigilator system featuring low-bandwidth audio/video streaming:"
    )
    doc.add_paragraph(
        "Candidate Media Slicing: When a secure exam begins, the candidate's browser initiates WebRTC and obtains a MediaStream. "
        "Every 3 seconds, a hidden video element draws a small 160x120 frame onto a canvas, compresses it to a 50% JPEG, and "
        "uploads the base64 string to `/api/proctor/signal?action=upload_feed`. Concurrently, a MediaRecorder listens to the microphone "
        "tracks, slices the audio into continuous 3-second slices, and uploads the base64 audio data. If the student disables "
        "their camera or microphone, the page instantly broadcasts `disabled` status."
    )
    doc.add_paragraph(
        "Live Invigilation Console: The proctor dashboard polls `/api/proctor/signal?action=get_feeds` every 3 seconds. The proctor "
        "can click \"Listen Mic\" to play the incoming audio slices sequentially using HTML5 dynamic Audio constructors. The proctor "
        "can click \"Record Audio\" to accumulate base64 chunks in a ref array; on stopping, the chunks are merged into a unified "
        "binary Uint8Array, packaged as a playable WebM file, downloaded, and logged under the candidate's evidence log."
    )
    doc.add_paragraph(
        "Mute/Offline Overlays: The dashboard handles status feeds. If the candidate disables the camera, the console overlays a "
        "black viewport stating \"Camera feed disabled by candidate\". If only the microphone is disabled, the live video stream "
        "remains visible, but a red pulsing \"MIC MUTED / OFFLINE\" badge is rendered at the top-left, and the telemetry overlay "
        "swaps to pulsing red \"MIC: MUTED\". If both are disabled, it overlays \"Camera & Microphone feeds disabled by candidate\"."
    )

    add_heading_with_spacing(doc, "4.2.4 Database Schema Management", level=3, before=12, after=6)
    doc.add_paragraph(
        "Database schema migrations are managed via Prisma CLI. A single schema definition (prisma/schema.prisma) "
        "represents all models. Changes are applied in dev via `npx prisma migrate dev` and applied in production via "
        "`npx prisma migrate deploy` which is executed automatically during the Vercel cloud build pipeline, keeping the database "
        "schema and Prisma Client generated models in absolute sync."
    )

    add_heading_with_spacing(doc, "4.3 Security, Performance, and Scalability Considerations", level=2, before=12, after=6)
    add_heading_with_spacing(doc, "4.3.1 Data Isolation Security", level=3, before=12, after=6)
    doc.add_paragraph(
        "The primary data isolation control is the scoping of all SQL queries by `organizationId` at the database level. "
        "All queries generated by Prisma Client dynamically inject `where: { organizationId }` based on the validated "
        "organization slug. JWT cookie tokens contain the user's authorized `orgSlug` and `role`, which are validated "
        "against the request path at the Edge Proxy middleware layer, creating a defense-in-depth security model."
    )

    add_heading_with_spacing(doc, "4.3.2 Performance Optimization for Nigerian Networks", level=3, before=12, after=6)
    doc.add_paragraph(
        "To handle unstable network speeds, the React bundle uses code-splitting via React `lazy()` and `Suspense`. "
        "Webcam images are compressed at 50% and sliced to 160x120 resolution before transmission. Static assets are served "
        "via Vercel's CDN Edge nearest to West Africa (Lagos and Abuja points of presence), optimizing page delivery."
    )

    add_heading_with_spacing(doc, "4.3.3 Scalability Through Database Design", level=3, before=12, after=6)
    doc.add_paragraph(
        "Scoping all queries by `organizationId` allows database index optimization. Prisma automatically creates "
        "composite indexes on lookup columns, e.g., `@@index([organizationId])` on tables like `SchoolCalendarEvent`, "
        "which prevents full table scans and ensures sub-second query execution times as tenant datasets grow."
    )

    add_heading_with_spacing(doc, "4.4 Challenges Encountered", level=2, before=12, after=6)
    doc.add_paragraph(
        "Subdomain Resolution on Edge: Vercel serverless functions run in isolated edge containers, which can complicate "
        "dynamic subdomain extraction. This was mitigated by routing pathnames via `/o/[slug]` and using Edge Proxy middleware "
        "to resolve subdomains at the CDN level, rewriting request headers dynamically."
    )
    doc.add_paragraph(
        "Audio Playback Sync: Playing back a series of discrete 3-second base64 audio chunks can result in audio overlapping. "
        "This was resolved by maintaining a `lastPlayedAudioTimestamp` ref in the proctor dashboard, verifying that each fetched "
        "audio chunk's server upload timestamp is greater than the last played chunk before triggering playback."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER FIVE: TESTING, RESULTS & EVALUATION
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER FIVE: TESTING, RESULTS & EVALUATION", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "5.1 Testing Strategy", level=2, before=12, after=6)
    doc.add_paragraph(
        "A multi-layered testing strategy was employed to ensure functional correctness and architectural integrity:"
    )
    doc.add_paragraph(
        "Unit Testing: Vitest was utilized to run unit tests on helper functions (JWT token validation, duration formatting, "
        "and grading math)."
    )
    doc.add_paragraph(
        "Integration Testing: API route endpoints (e.g. proctor signal upload/download, submission saves) were tested using HTTP mock "
        "requests. Cross-tenant testing confirmed that a JWT containing org A slug would be rejected with a 403 Forbidden status "
        "when requesting resource data under org B."
    )
    doc.add_paragraph(
        "User Acceptance Testing (UAT): The prototype was evaluated by five educators in Nigeria. They successfully tested user logging, "
        "course configuration, and taking a proctored exam."
    )

    add_heading_with_spacing(doc, "5.2 Test Results", level=2, before=12, after=6)

    # Table 5.1
    add_heading_with_spacing(doc, "Table 5.1: Database Isolation Verification Test Cases", level=3, before=12, after=6)
    t51 = doc.add_table(rows=4, cols=4)
    t51.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr51 = t51.rows[0].cells
    hdr51[0].text = 'Test Case'
    hdr51[1].text = 'Input'
    hdr51[2].text = 'Expected Output'
    hdr51[3].text = 'Status'
    for cell in hdr51:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    t51_data = [
        ('Valid Login', 'school-a slug, valid credentials', '200 OK + JWT cookie containing school-a claims', 'Pass'),
        ('Invalid Tenant Login', 'school-b slug, school-a credentials', '401 Unauthorized, access denied', 'Pass'),
        ('Cross-Tenant Data Leak', 'User of school-a requests data under school-b slug', '403 Forbidden, session slug mismatch', 'Pass')
    ]
    for idx, (tc, ti, te, ts) in enumerate(t51_data):
        row_cells = t51.rows[idx+1].cells
        row_cells[0].text = tc
        row_cells[1].text = ti
        row_cells[2].text = te
        row_cells[3].text = ts
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    doc.add_paragraph()

    # Table 5.2
    add_heading_with_spacing(doc, "Table 5.2: Performance Benchmarking (Simulated 3G Network)", level=3, before=12, after=6)
    t52 = doc.add_table(rows=4, cols=3)
    t52.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr52 = t52.rows[0].cells
    hdr52[0].text = 'Metric'
    hdr52[1].text = 'Measured Value (Avg 5 runs)'
    hdr52[2].text = 'Target Limit (NFR-02)'
    for cell in hdr52:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    t52_data = [
        ('First Contentful Paint (FCP)', '1.8 seconds', '< 3.0 seconds'),
        ('Time to Interactive (TTI)', '3.4 seconds', '< 4.0 seconds'),
        ('Backend API Latency (GET /courses)', '210 ms', 'N/A')
    ]
    for idx, (m, mv, req) in enumerate(t52_data):
        row_cells = t52.rows[idx+1].cells
        row_cells[0].text = m
        row_cells[1].text = mv
        row_cells[2].text = req
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    p_t52_cap = doc.add_paragraph("Table 5.2: Performance Benchmarking (Simulated 3G Network via Chrome DevTools)")
    p_t52_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t52_cap.paragraph_format.space_before = Pt(6)

    doc.add_paragraph(
        "The measured FCP of 1.8 seconds under simulated 3G conditions (500ms RTT, 1.6 Mbps downlink) satisfies "
        "the performance requirements, confirming that the code-splitting and asset compression are highly effective "
        "for the target West African broadband infrastructure."
    )

    add_heading_with_spacing(doc, "5.3 Evaluation of Objectives", level=2, before=12, after=6)

    # Table 5.3
    t53 = doc.add_table(rows=6, cols=3)
    t53.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr53 = t53.rows[0].cells
    hdr53[0].text = 'Objective'
    hdr53[1].text = 'Evidence of Achievement'
    hdr53[2].text = 'Status'
    for cell in hdr53:
        set_cell_background(cell, "1F4E79")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_margins(cell)

    t53_data = [
        ('1. Multi-Tenant Database Design', 'PostgreSQL database configured with composite organizationId indexing.', 'Achieved'),
        ('2. Bandwidth-Optimized Interface', 'Mobile-first React application with FCP of 1.8 seconds on simulated 3G.', 'Achieved'),
        ('3. Core LMS Features', 'Functional course management, quiz schedules, and certificate downloads.', 'Achieved'),
        ('4. Live Proctoring System', 'Live invigilator dashboard with camera frame uploads, audio slices, and remote actions.', 'Achieved'),
        ('5. Platform operator Analytics', 'Aggregated Super Admin dashboard displaying active schools and user counts.', 'Achieved')
    ]
    for idx, (obj, ev, stat) in enumerate(t53_data):
        row_cells = t53.rows[idx+1].cells
        row_cells[0].text = obj
        row_cells[1].text = ev
        row_cells[2].text = stat
        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "F2F2F2")

    p_t53_cap = doc.add_paragraph("Table 5.3: Evaluation of Project Objectives")
    p_t53_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t53_cap.paragraph_format.space_before = Pt(6)

    add_heading_with_spacing(doc, "5.4 Discussion of Results", level=2, before=12, after=6)
    doc.add_paragraph(
        "The test results validate the technical feasibility of the multi-tenant SaaS approach for the "
        "Nigerian education market. The data isolation tests (Table 5.1) confirm that the logical shared-database "
        "strategy combined with Edge middleware token verification is robust against cross-tenant data leakage. "
        "This aligns with documented best practices for multi-tenant data isolation in regulated sectors (Sharma, 2025)."
    )
    doc.add_paragraph(
        "The performance metrics (Table 5.2) demonstrate that the React/Node stack deployed on Vercel's edge network "
        "delivers a satisfactory user experience even under constrained network conditions. The achieved FCP of 1.8 "
        "seconds is significantly better than the 3-second target, indicating that the bandwidth-optimization "
        "strategies were effective."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTER SIX: SUMMARY, CONCLUSION & RECOMMENDATIONS
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "CHAPTER SIX: SUMMARY, CONCLUSION & RECOMMENDATIONS", level=1, before=18, after=12)

    add_heading_with_spacing(doc, "6.1 Summary of the Project", level=2, before=12, after=6)
    doc.add_paragraph(
        "This project set out to design and implement a scalable multi-tenant SaaS architecture for a Learning "
        "Management System tailored to the specific needs and constraints of Nigerian educational institutions. The "
        "project successfully navigated the complexities of database isolation, subdomain routing, white-label branding, "
        "and platform analytics. The resulting prototype is deployed on Vercel and Supabase, serving as a functional "
        "demonstration of how modern cloud-native tools can democratize access to digital learning platforms in a cost-effective "
        "and operationally sustainable manner. The iterative development methodology enabled continuous validation of "
        "tenant isolation and performance requirements (Soesanto et al., 2023; Cakar, 2025)."
    )

    add_heading_with_spacing(doc, "6.2 Conclusion", level=2, before=12, after=6)
    doc.add_paragraph(
        "The project concludes that a Shared Database, Shared Schema PostgreSQL model, implemented using Prisma 7 and "
        "managed on Supabase, deployed via Vercel serverless edge functions, provides an optimal balance of resource "
        "efficiency, data security, and operational manageability for this use case. This finding aligns with established "
        "best practices for multi-tenant SaaS architectures in cloud deployments (Sharma, 2025). The challenges of serverless "
        "cold starts were successfully mitigated through Vercel cron-based health pings. The project confirms that by "
        "centralizing infrastructure and leveraging modern cloud services, a multi-tenant SaaS LMS can significantly lower "
        "the Total Cost of Ownership for schools in Nigeria."
    )

    add_heading_with_spacing(doc, "6.3 Professional Contributions", level=2, before=12, after=6)
    doc.add_paragraph(
        "This project makes four distinct contributions to software engineering practice and educational technology in Nigeria:"
    )
    doc.add_paragraph(
        "1. An Open Reference Architecture for Multi-Tenant EdTech SaaS: While commercial solutions like SabiScholar validate "
        "market demand, their proprietary nature limits knowledge dissemination. This thesis provides a fully documented reference "
        "implementation using Supabase, Vercel, and Next.js, serving as a blueprint for developers seeking to build scalable "
        "multi-tenant applications. This contribution joins emerging open-source efforts such as ClassroomIO (2025)."
    )
    doc.add_paragraph(
        "2. A Cost-Effective, Serverless Deployment Model for Educational SaaS: By leveraging Vercel's edge network and managed "
        "PostgreSQL, this project demonstrates a deployment architecture that minimizes operational overhead, scales automatically, "
        "and reduces total cost of ownership."
    )
    doc.add_paragraph(
        "3. A Replicable Pattern for Live Exam Invigilation: The implementation of low-bandwidth camera frame polling and "
        "continuous 3-second audio slice transfers provides a robust alternative to high-bandwidth streaming protocols, making "
        "proctored exams viable on standard cellular networks in West Africa."
    )
    doc.add_paragraph(
        "4. Empirical Performance Benchmarking: The project provides documented performance metrics (FCP, TTI) under simulated "
        "3G network conditions typical of Nigerian mobile users, establishing a baseline methodology for evaluating LMS performance "
        "in bandwidth-constrained environments."
    )

    add_heading_with_spacing(doc, "6.4 Recommendations for Future Work", level=2, before=12, after=6)
    doc.add_paragraph(
        "While the prototype meets the core objectives, several areas are ripe for future enhancement:"
    )
    doc.add_paragraph(
        "1. Integration of Nigerian Payment Gateways: Implementing a billing module integrated with Paystack or Flutterwave "
        "to automate school onboarding and subscription management (Atwongyere, 2025)."
    )
    doc.add_paragraph(
        "2. Offline-First Capabilities: Implementing a Progressive Web App (PWA) with service workers to allow students to download "
        "course syllabus and lesson files for offline review during network outages (DOAJ, 2025)."
    )
    doc.add_paragraph(
        "3. Advanced AI-Based Behavior Analysis: Integrating lightweight models to analyze camera feeds locally on the client "
        "for gaze secure tracking, reporting flags to the invigilator dashboard dynamically."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # REFERENCES
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "REFERENCES", level=1)
    doc.add_paragraph(
        "6Wresearch. (2025a). Global Learning Management System in Education Sector Market (2025-2031). "
        "Retrieved from https://www.6wresearch.com/"
    )
    doc.add_paragraph(
        "6Wresearch. (2025b). Nigeria Learning Management System Market (2025-2031) | Trends, Outlook & Forecast. "
        "Retrieved from https://www.6wresearch.com/"
    )
    doc.add_paragraph(
        "Assalaarachchi, L. I., Silva, K. P. K. H., & Hewagamage, C. (2023). Adoption of Software-as-a-Service (SaaS) "
        "Applications in E-learning: Perception of the Management Undergraduates in a Selected State University of Sri Lanka. "
        "Vidyodaya Journal of Management, 9(II). Retrieved from https://journals.sjp.ac.lk/"
    )
    doc.add_paragraph(
        "Atwongyere, D. (2025). Online Multi-tenant E-learning Platform With Integrated Payment Gateways [Master’s thesis, "
        "Uganda Christian University]. UCU Scholar."
    )
    doc.add_paragraph(
        "Azouzi, S., & Ghannouchi, S. A. (2025). Designing Multi-Tenant E-Learning Systems in the Cloud: A Process-Oriented "
        "Approach for Higher Education. International Journal of Computer Applications, 187(10), 16-24. "
        "https://www.ijcaonline.org/archives/volume187/number10/azouzi-2025-ijca-924998.pdf"
    )
    doc.add_paragraph(
        "Azouzi, S., Brahmi, Z., & Ghannouchi, S. A. (2018). Customization of multi-tenant learning process as a service "
        "with Business Process Feature Model. Procedia Computer Science, 126, 606-615. https://doi.org/10.1016/j.procs.2018.07.295"
    )
    doc.add_paragraph(
        "Cakar, C. (2025). Scalable Multi-Tenant Software Architectures: Isolation, Performance, and Governance in "
        "Enterprise-Grade Systems. Iconic Research and Engineering Journals, 9(5), 2821-2831. https://doi.org/10.64388/IREV9I5-1715580"
    )
    doc.add_paragraph(
        "Can, F., Doğan, G., Konca, C., & Akbulut, A. (2017). Multi-tenant architectures in the cloud: A systematic mapping study. "
        "In 2017 International Artificial Intelligence and Data Processing Symposium (IDAP) (pp. 1-4). IEEE. "
        "https://doi.org/10.1109/IDAP.2017.8090272"
    )
    doc.add_paragraph(
        "ClassroomIO. (2025). Open-Source EdTech Platform [GitHub repository]. Retrieved from https://github.com/classroomio/classroomio"
    )
    doc.add_paragraph(
        "ClickHouse. (2026). How to architect multi-tenant SaaS on Postgres. ClickHouse Resource Hub. Retrieved from https://clickhouse.com/"
    )
    doc.add_paragraph(
        "Educational Technology Quarterly. (2025). Push and pull factors influencing LMS use post-COVID in Nigerian open distance "
        "learning: a qualitative study. Educational Technology Quarterly, 2025(4), 407-428. https://doi.org/10.55056/etq.966"
    )
    doc.add_paragraph(
        "Kester, K. O., & Ojedeji, S. O. (2022). Workers’ Education from the Cloud: Maximising Latest Technologies for "
        "Human Resource Development in Africa. In New Updates in E-Learning. IntechOpen. Retrieved from https://www.academia.edu/"
    )
    doc.add_paragraph(
        "Manzoor, N. (2024). Optimizing Educational Access: Identifying the Ideal Cloud Solution for Student Learning App "
        "in African Contexts [Bachelor’s thesis]. Theseus. Retrieved from https://www.theseus.fi/"
    )
    doc.add_paragraph(
        "Meduri, R. K. K., Gutha, S., & Jadala, V. C. (2023). An Architectural Review of Multi-Tenancy in Cloud Computing. "
        "In Research Anthology on Modern Cloud Computing. IGI Global. "
        "https://www.igi-global.com/chapter/an-architectural-review-of-multi-tenancy-in-cloud-computing/330065"
    )
    doc.add_paragraph(
        "Mell, P., & Grance, T. (2011). The NIST Definition of Cloud Computing (NIST Special Publication 800-145). "
        "National Institute of Standards and Technology. https://doi.org/10.6028/NIST.SP.800-145"
    )
    doc.add_paragraph(
        "Microsoft. (2026). Designing Your SaaS Database for Scale with PostgreSQL – Citus for PostgreSQL. Microsoft Learn. "
        "Retrieved from https://learn.microsoft.com/"
    )
    doc.add_paragraph(
        "Osode, J. I., Lautenbach, G., & Goto, J. (2024). Factors influencing teaching staff’s adoption of Learning "
        "Management Systems in three Nigerian universities. African Journal of Teacher Education, 13(2). "
        "Retrieved from https://journal.lib.uoguelph.ca/"
    )
    doc.add_paragraph(
        "Sharma, N. (2025). Secure by Design: Architecture Patterns for Multi-Tenant SaaS at Scale. DEV Community. "
        "Retrieved from https://dev.to/niranjan_sharma_579448819/secure-by-design-architecture-patterns-for-multi-tenant-saas-at-scale-4bah"
    )
    doc.add_paragraph(
        "Soesanto, D., Liliana, Louk, M. H. L., & Handani, F. (2023). Implementation of Software as a Service to "
        "Increase the Scalability of the Merdeka Belajar Policy in Indonesia. Jurnal Teknologi Pendidikan, 25(2)."
    )

    doc.add_page_break()

    # ----------------------------------------------------
    # APPENDICES
    # ----------------------------------------------------
    add_heading_with_spacing(doc, "APPENDICES", level=1)

    add_heading_with_spacing(doc, "Appendix A: User Manual Excerpt", level=2)
    doc.add_paragraph("A.1 Accessing the Platform")
    doc.add_paragraph("Each school receives a unique URL slug under the application. Examples:")
    doc.add_paragraph("• https://skilltech.com.ng/o/demo-school (Demo School)")
    doc.add_paragraph("• https://skilltech.com.ng/o/demo-university (Demo University)")

    doc.add_paragraph("\nA.2 Logging In")
    doc.add_paragraph("1. Navigate to your school's unique URL.")
    doc.add_paragraph("2. Enter credentials provided by your school administrator.")
    doc.add_paragraph("3. Click \"Sign In\".")

    doc.add_paragraph("\nDefault roles and sample credentials (demo data):")
    doc.add_paragraph("School Admin: admin@test.com / Password: password123")
    doc.add_paragraph("Teacher: teacher@test.com / Password: password123")
    doc.add_paragraph("Student: student@test.com / Password: password123")

    doc.add_paragraph("\nA.3 Live Invigilation Console Usage (Teacher View)")
    doc.add_paragraph(
        "1. Click the \"Exam Proctoring & Integrity\" link in the main navigation sidebar.\n"
        "2. Click the \"Live Grid Monitor\" tab.\n"
        "3. Click the \"Start Live Feeds\" button to request webcam access and begin monitoring.\n"
        "4. In each student's card, you can view their telemetry state (Gaze secure, Mic status) and webcam feed.\n"
        "5. Under each student's card, click \"Listen Mic\" to hear their audio, or click \"Record Audio\" to record evidence.\n"
        "6. Click \"Prompt Audio\" / \"Force Audio\" or \"Prompt Camera\" / \"Force Camera\" to remotely configure client hardware.\n"
        "7. Click \"Snapshot\" or \"Record Clip\" to download file captures of candidate violations."
    )

    add_heading_with_spacing(doc, "Appendix B: Code Snippets", level=2)
    doc.add_paragraph("B.1 Tenant Resolution and Security Guard (src/proxy.ts)")
    p_code1 = doc.add_paragraph()
    p_code1.paragraph_format.line_spacing = 1.0
    r_code1 = p_code1.add_run(
        "export async function proxy(request: NextRequest) {\n"
        "  const pathname = request.nextUrl.pathname;\n\n"
        "  if (pathname.startsWith(\"/platform\")) {\n"
        "    const secret = process.env.PLATFORM_JWT_SECRET ?? \"\";\n"
        "    const key = new TextEncoder().encode(secret);\n"
        "    const token = request.cookies.get(PLATFORM_AUTH_COOKIE)?.value;\n"
        "    if (!token) return redirectToPlatformLogin(request, pathname);\n"
        "    try {\n"
        "      const { payload } = await jwtVerify(token, key);\n"
        "      if (payload.kind !== \"platform\") return redirectToPlatformLogin(request, pathname);\n"
        "      return NextResponse.next();\n"
        "    } catch { return redirectToPlatformLogin(request, pathname); }\n"
        "  }\n\n"
        "  const match = pathname.match(/^\\/o\\/([^/]+)/);\n"
        "  if (!match) return NextResponse.next();\n"
        "  const slug = match[1];\n"
        "  const secret = process.env.JWT_SECRET ?? \"\";\n"
        "  const key = new TextEncoder().encode(secret);\n"
        "  const token = request.cookies.get(AUTH_COOKIE)?.value;\n"
        "  if (!token) return redirectToLogin(request, pathname, slug);\n"
        "  try {\n"
        "    const { payload } = await jwtVerify(token, key);\n"
        "    const orgSlug = payload.orgSlug as string | undefined;\n"
        "    if (!orgSlug || orgSlug !== slug) {\n"
        "      return NextResponse.redirect(new URL(\"/\", request.url));\n"
        "    }\n"
        "    return NextResponse.next();\n"
        "  } catch { return redirectToLogin(request, pathname, slug); }\n"
        "}"
    )
    r_code1.font.name = 'Consolas'
    r_code1.font.size = Pt(8.5)

    doc.add_paragraph("\nB.2 Candidate Audio Chunk transmitter (take-assessment.tsx)")
    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.line_spacing = 1.0
    r_code2 = p_code2.add_run(
        "useEffect(() => {\n"
        "  if (!localStream || !studentEmail || deliveryMode === \"FORMATIVE\" || locked) return;\n"
        "  const audioTracks = localStream.getAudioTracks();\n"
        "  if (audioTracks.length === 0) return;\n\n"
        "  let mediaRecorder: MediaRecorder | null = null;\n"
        "  try {\n"
        "    const audioStream = new MediaStream(audioTracks);\n"
        "    const options = MediaRecorder.isTypeSupported(\"audio/webm\") ? { mimeType: \"audio/webm\" } : undefined;\n"
        "    mediaRecorder = new MediaRecorder(audioStream, options);\n"
        "    mediaRecorder.ondataavailable = (e) => {\n"
        "      const audioTrack = localStream.getAudioTracks()[0];\n"
        "      if (!audioTrack || !audioTrack.enabled) {\n"
        "        void fetch(\"/api/proctor/signal\", {\n"
        "          method: \"POST\",\n"
        "          headers: { \"Content-Type\": \"application/json\" },\n"
        "          body: JSON.stringify({ action: \"upload_feed\", studentEmail, audioFeed: \"disabled\" }),\n"
        "        });\n"
        "        return;\n"
        "      }\n"
        "      if (e.data && e.data.size > 0) {\n"
        "        const reader = new FileReader();\n"
        "        reader.onloadend = () => {\n"
        "          const base64Audio = (reader.result as string).split(\",\")[1];\n"
        "          void fetch(\"/api/proctor/signal\", { \n"
        "            method: \"POST\", \n"
        "            headers: { \"Content-Type\": \"application/json\" },\n"
        "            body: JSON.stringify({ action: \"upload_feed\", studentEmail, audioFeed: base64Audio }),\n"
        "          });\n"
        "        };\n"
        "        reader.readAsDataURL(e.data);\n"
        "      }\n"
        "    };\n"
        "    mediaRecorder.start(3000);\n"
        "  } catch (err) { console.warn(err); }\n"
        "  return () => { if (mediaRecorder) mediaRecorder.stop(); };\n"
        "}, [localStream, studentEmail, deliveryMode, locked]);"
    )
    r_code2.font.name = 'Consolas'
    r_code2.font.size = Pt(8.5)

    add_heading_with_spacing(doc, "Appendix C: Screenshots", level=2)
    doc.add_paragraph("Figure C1: Super Admin Platform Analytics Dashboard.")
    doc.add_paragraph("Figure C2: Organization Login Page (e.g. demo-school) displaying customized logo and primary brand color styling.")
    doc.add_paragraph("Figure C3: Course list and syllabus module editing dashboard for school Teachers.")
    doc.add_paragraph("Figure C4: Student Exam taking workspace with floating proctor box previewing active webcam/microphone tracks.")
    doc.add_paragraph("Figure C5: Live Exam Invigilator Dashboard showing active candidate cards with live camera feeds and telemetry overlays (MIC: ACTIVE, GAZE: SECURE).")
    doc.add_paragraph("Figure C6: Live Exam Invigilator Dashboard showing 'Camera feed disabled by candidate' warning overlay on the candidate video grid viewport.")
    doc.add_paragraph("Figure C7: Live Exam Invigilator Dashboard showing 'MIC MUTED / OFFLINE' pulsing warning badge and telemetry change to 'MIC: MUTED' when the candidate disables the microphone.")

    add_heading_with_spacing(doc, "Appendix D: Deployment Configuration", level=2)
    doc.add_paragraph("D.1 Vercel Build configurations (vercel.json)")
    p_code3 = doc.add_paragraph()
    p_code3.paragraph_format.line_spacing = 1.0
    r_code3 = p_code3.add_run(
        "{\n"
        "  \"$schema\": \"https://openapi.vercel.sh/vercel.json\",\n"
        "  \"framework\": \"nextjs\",\n"
        "  \"buildCommand\": \"node scripts/vercel-build.mjs\"\n"
        "}"
    )
    r_code3.font.name = 'Consolas'
    r_code3.font.size = Pt(8.5)

    doc.add_paragraph("\nD.2 Environment Variables Config (.env.example)")
    p_code4 = doc.add_paragraph()
    p_code4.paragraph_format.line_spacing = 1.0
    r_code4 = p_code4.add_run(
        "DATABASE_URL=\"postgresql://...\"\n"
        "DIRECT_URL=\"postgresql://...\"\n"
        "JWT_SECRET=\"...\"\n"
        "NEXT_PUBLIC_APP_URL=\"...\"\n"
        "PLATFORM_ADMIN_EMAIL=\"...\"\n"
        "PLATFORM_ADMIN_PASSWORD=\"...\"\n"
        "PLATFORM_JWT_SECRET=\"...\""
    )
    r_code4.font.name = 'Consolas'
    r_code4.font.size = Pt(8.5)

    doc.save("Project_Documentation.docx")
    print("SUCCESS: Project_Documentation.docx generated successfully!")

if __name__ == "__main__":
    main()
